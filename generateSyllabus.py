from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_UNDERLINE
import openpyxl, os
import pandas as pd

document = Document()
section = document.sections[0]
header = section.header

f = open( "q&a.txt", "r")
g = f.read()
answers = [s.replace("[", "").replace("]", "") for s in g.split('"')]

# Load schedule excel
df = pd.read_excel('C:/Users/mahya/Desktop/Capstone/courseorganizer-main/Scripts/test.xlsx', sheet_name='Sheet1')

def generate():

# Template of the syllabus
    paragraph = header.paragraphs[0]
    paragraph.text = "Course Syllabus\t\t" + answers[23]
    paragraph.style = document.styles["Header"]

    title = document.add_paragraph(answers[1])
    title_format = title.paragraph_format
    title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title2 = document.add_paragraph(answers[2])
    title2_format = title2.paragraph_format
    title2_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title3 = document.add_paragraph(answers[13])
    title3_format = title3.paragraph_format
    title3_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    records = [
        ('Class Location: ' + answers[15]), ('E-mail: ' + answers[11]),
        ('Office location: ' + answers[21]), ('Office hours: ' + answers[19]),
    
    ]

    if (answers[14].upper() == 'YES'):
        records.append(('Phone: ' + answers[9]))
    if (answers[22].upper() == 'YES'):
        records.append(('Discussion seminar time: ' + answers[49]), ('Discussion Seminar location: ' + answers[51]))
    if (answers[35].upper() == 'YES'):
        records.append(('Virtual office hours(Zoom)\nMeeting ID: '+ answers[37] + '\nPasscode: ' + answers[39]))

    num_rows = (len(records) + 1) // 2
    table = document.add_table(rows=num_rows, cols=2, style = "Table Grid")
    header_row = table.rows[0].cells
    header_row[0].text = 'Instructor: ' + answers[5]
    header_row[1].text = 'Class Time: ' + answers[17]
#for x, y in records:
#    row_cells = table.add_row().cells
#    row_cells[0].text = x
#    row_cells[1].text = y

    for i in range(len(records)):
        row_idx = (i + 1) // 2  # add 1 to account for header row
        col_idx = (i + 1) % 2
        if col_idx == 0:
            col_idx = 2
        cell = table.rows[row_idx].cells[col_idx - 1]
        cell.text = records[i]

    paragraph2 = document.add_paragraph('\n')

    if (answers[77].upper() == 'YES'):
        run = paragraph2.add_run('Prerequisites')
        run.underline = True
        run.underline = WD_UNDERLINE.SINGLE
        paragraph2.add_run('\n' + answers[79])

    paragraph2.add_run('\n')
    paragraph2.add_run('\nCourse Description').underline = WD_UNDERLINE.SINGLE
    paragraph2.add_run('\n' + answers[31])

    if (answers[81].upper() == 'YES'):
        paragraph2.add_run('\n')
        paragraph2.add_run('\nCourse Materials').underline = WD_UNDERLINE.SINGLE
        paragraph2.add_run('\n' 'Textbook: ' + answers[83])
        paragraph2.add_run('\n' 'ISBN: ' + answers[85])

    if (answers[87].upper() == 'YES'):
        paragraph2.add_run('\n')
        paragraph2.add_run('\nLearning Objectives').underline = WD_UNDERLINE.SINGLE
        paragraph2.add_run('\n' + answers[89])

    if (answers[43].upper() == 'YES'):
        paragraph2.add_run('\n')
        paragraph2.add_run('\nLab Policy').underline = WD_UNDERLINE.SINGLE
        paragraph2.add_run('\n' + answers[45])

    if (answers[59].upper() == 'YES'):
        paragraph2.add_run('\n')
        paragraph2.add_run('\nAssignments').underline = WD_UNDERLINE.SINGLE
        paragraph2.add_run('\n' + answers[61])

    paragraph2.add_run('\n')
    paragraph2.add_run('\nExpectations').underline = WD_UNDERLINE.SINGLE
    paragraph2.add_run('\n' + answers[91])

    if (answers[63].upper() == 'YES'):
        paragraph2.add_run('\n')
        paragraph2.add_run('\nQuiz').underline = WD_UNDERLINE.SINGLE
        paragraph2.add_run('\n' + answers[65])

    if (answers[69].upper() == 'YES'):
        paragraph2.add_run('\n')
        paragraph2.add_run('\nExam').underline = WD_UNDERLINE.SINGLE
        paragraph2.add_run('\n' + answers[71])

    if (answers[47].upper() == 'YES'):
        paragraph2.add_run('\n')
        paragraph2.add_run('\nDiscussion Policy').underline = WD_UNDERLINE.SINGLE
        paragraph2.add_run('\n' + answers[49])


    paragraph2.add_run('\n')
    paragraph2.add_run('\nAttendance').underline = WD_UNDERLINE.SINGLE
    paragraph2.add_run('\n' + answers[29])

    paragraph2.add_run('\n')
    paragraph2.add_run('\nGrading').underline = WD_UNDERLINE.SINGLE
    paragraph2.add_run('\n' + answers[95])

    paragraph2.add_run('\n')
    paragraph2.add_run('\nDisability Services').underline = WD_UNDERLINE.SINGLE
    paragraph2.add_run('\n' + answers[27])

    paragraph2.add_run('\n')
    paragraph2.add_run('\nHonor Code').underline = WD_UNDERLINE.SINGLE
    paragraph2.add_run('\n' + answers[25])

    if (answers[97].upper() == 'YES'):
        paragraph2.add_run('\n')
        paragraph2.add_run('\nOnline Resources').underline = WD_UNDERLINE.SINGLE
        paragraph2.add_run('\n' + answers[99])

    if (answers[101].upper() == 'YES'):
        paragraph2.add_run('\n')
        paragraph2.add_run('\nExtra Credit').underline = WD_UNDERLINE.SINGLE
        paragraph2.add_run('\n' + answers[103])

    paragraph2.add_run('\n')
    paragraph2.add_run('\nFinal Exam').underline = WD_UNDERLINE.SINGLE
    paragraph2.add_run('\n' + answers[105])

    paragraph2.add_run('\n')
    paragraph2.add_run('\nInclement Weather').underline = WD_UNDERLINE.SINGLE
    paragraph2.add_run('\n' + answers[33])

    paragraph2.add_run('\n')
    paragraph2.add_run('\nWithdrawals').underline = WD_UNDERLINE.SINGLE
    paragraph2.add_run('\n' + answers[107])

    if (answers[27].upper() == 'YES'):
        paragraph2.add_run('\n')
        paragraph2.add_run('\nSyllabus Update').underline = WD_UNDERLINE.SINGLE
        paragraph2.add_run('\nThis syllabus can be updated at any point in time in the semester')

    paragraph2.add_run('\n')
    paragraph2.add_run('\nLecture Schedule').underline = WD_UNDERLINE.SINGLE

#schedule table
    df.fillna('', inplace=True)
    table2 = document.add_table(rows=1, cols=len(df.columns), style = "Table Grid")

# Add header row to the table
    hdr_cells = table2.rows[0].cells
    for i in range(len(df.columns)):
        hdr_cells[i].text = df.columns[i]

# Add data to the table
    for index, row in df.iterrows():
        row_cells = table2.add_row().cells
        for i in range(len(df.columns)):
            row_cells[i].text = str(row[i])

    document.save('syllabus-generated.docx')

generate()